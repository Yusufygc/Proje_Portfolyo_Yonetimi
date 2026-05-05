"""infrastructure/export/docx_ats_exporter.py — ATS uyumlu DOCX CV üretici."""

import io
import logging

from services.interfaces.i_exporter import IExporter
from services.dto.export_config import ExportConfig

logger = logging.getLogger(__name__)


class DOCXATSExporter(IExporter):
    """ATS uyumlu DOCX formatında CV üretici (python-docx kullanır)."""

    def get_format_name(self) -> str:
        return "ATS Uyumlu CV (DOCX)"

    def get_file_extension(self) -> str:
        return ".docx"

    def export(
        self,
        personal_info,
        projects: list,
        skills: list,
        education: list,
        experience: list,
        certificates: list,
        config: ExportConfig,
    ) -> bytes:
        """ATS uyumlu DOCX üretir ve bytes döner."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError(
                "python-docx kütüphanesi bulunamadı. "
                "Lütfen 'pip install python-docx' komutunu çalıştırın."
            )

        doc = Document()
        lang = config.language
        full_name = getattr(personal_info, "full_name", "CV") or "CV"

        # Margin: 1 inch her yönden
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        def add_heading(text: str, level: int = 1):
            para = doc.add_paragraph()
            run = para.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)
            # İnce çizgi: DOCX'te bottom border ile simüle et
            return para

        def add_body(text: str, bold: bool = False, font_size: int = 11):
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.name = "Calibri"
            para.paragraph_format.space_after = Pt(4)
            return para

        # 1. İletişim Bloğu
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(full_name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_val = getattr(personal_info, "title", "") or ""
        if title_val:
            title_para = doc.add_paragraph(title_val)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].font.size = Pt(12)

        # Tek satır iletişim
        parts = []
        email = getattr(personal_info, "email", None)
        phone = getattr(personal_info, "phone", None)
        github = getattr(personal_info, "github_url", None)
        linkedin = getattr(personal_info, "linkedin_url", None)

        if email:
            parts.append(email)
        if phone:
            parts.append(phone)
        if github:
            parts.append(github.replace("https://", ""))
        if linkedin:
            parts.append(linkedin.replace("https://", ""))

        if parts:
            contact_para = doc.add_paragraph(" | ".join(parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # Boşluk

        # 2. Professional Summary
        bio = getattr(personal_info, "bio", "") or ""
        if bio:
            heading_text = "PROFESSIONAL SUMMARY" if lang == "en" else "PROFESYONEl ÖZET"
            add_heading(heading_text)
            summary = bio
            if config.target_role:
                summary = (
                    f"Results-driven {config.target_role} with expertise in "
                    f"software development. {bio}"
                )
            add_body(summary)

        # 3. Technical Skills
        if skills:
            heading_text = "TECHNICAL SKILLS" if lang == "en" else "TEKNİK BECERİLER"
            add_heading(heading_text)
            from collections import defaultdict
            grouped: dict = defaultdict(list)
            for s in skills:
                cat = (getattr(s, "category", "") or "OTHER").upper()
                grouped[cat].append(s.name)

            labels = {
                "LANGUAGE":  "Languages",
                "FRAMEWORK": "Frameworks",
                "DATABASE":  "Databases",
                "TOOL":      "Tools",
                "DEVOPS":    "DevOps",
            }
            order = ["LANGUAGE", "FRAMEWORK", "DATABASE", "TOOL", "DEVOPS"]
            cats = [c for c in order if c in grouped] + [c for c in grouped if c not in order]
            for cat in cats:
                label = labels.get(cat, cat.title())
                para = doc.add_paragraph()
                run_bold = para.add_run(f"{label}: ")
                run_bold.bold = True
                run_bold.font.size = Pt(11)
                run_normal = para.add_run(", ".join(grouped[cat]))
                run_normal.font.size = Pt(11)
                para.paragraph_format.space_after = Pt(4)

        # 4. Projects
        selected = (
            [p for p in projects if p.id in config.selected_project_ids]
            if config.selected_project_ids else projects
        )
        if selected:
            heading_text = "PROJECTS" if lang == "en" else "PROJELER"
            add_heading(heading_text)
            for proj in selected:
                add_body(proj.title, bold=True)
                desc = getattr(proj, "full_description", "") or ""
                if desc:
                    add_body(desc)
                tags = getattr(proj, "tags", [])
                if tags:
                    tech_str = ", ".join(t.tag_name for t in tags)
                    add_body(f"Technologies: {tech_str}", font_size=10)

        # 5. Experience
        if experience:
            heading_text = "EXPERIENCE" if lang == "en" else "DENEYİM"
            add_heading(heading_text)
            for exp in sorted(experience, key=lambda x: x.start_date or "", reverse=True):
                add_body(exp.company, bold=True)
                date_str = exp.start_date or ""
                if getattr(exp, "is_current", False):
                    date_str += " - Present"
                elif getattr(exp, "end_date", None):
                    date_str += f" - {exp.end_date}"
                add_body(f"{exp.position} | {date_str}", font_size=10)
                desc = getattr(exp, "description", "") or ""
                if desc:
                    for line in desc.split("\n"):
                        line = line.strip()
                        if line:
                            add_body(f"• {line}")

        # 6. Education
        if education:
            heading_text = "EDUCATION" if lang == "en" else "EĞİTİM"
            add_heading(heading_text)
            for edu in education:
                add_body(edu.institution, bold=True)
                degree = edu.degree
                field = getattr(edu, "field", "") or ""
                if field:
                    degree += f", {field}"
                date_str = edu.start_date or ""
                if getattr(edu, "end_date", None):
                    date_str += f" - {edu.end_date}"
                add_body(f"{degree} | {date_str}", font_size=10)

        # 7. Certifications
        if certificates:
            heading_text = "CERTIFICATIONS" if lang == "en" else "SERTİFİKALAR"
            add_heading(heading_text)
            for cert in certificates:
                add_body(f"{cert.name} - {cert.issuer} ({cert.date or ''})")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
